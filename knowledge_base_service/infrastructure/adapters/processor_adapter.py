"""
Document processor adapter
"""
import io
from typing import List, Optional

from docx import Document as DocxDocument
from pypdf import PdfReader

from config import get_settings
from domain.entities.document import Document, DocumentChunk
from domain.ports.processor_port import IDocumentProcessor
from infrastructure.adapters.embedding_adapter import EmbeddingService

settings = get_settings()


class DocumentProcessor(IDocumentProcessor):
    """Document processor implementation"""
    
    def __init__(self):
        self.embedding_service = EmbeddingService()
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    async def process_document(
        self,
        filename: str,
        content: bytes,
        content_type: Optional[str] = None
    ) -> Document:
        """Process document and extract chunks"""
        # Extract text based on file type
        if filename.endswith('.pdf'):
            text = self._extract_from_pdf(content)
        elif filename.endswith('.docx'):
            text = self._extract_from_docx(content)
        elif filename.endswith('.txt'):
            text = content.decode('utf-8')
        elif filename.endswith(('.md', '.markdown')):
            text = self._extract_from_markdown(content)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Create document
        document = Document(
            filename=filename,
            content_type=content_type,
            file_size=len(content),
            content=text
        )
        
        # Chunk document
        chunks = await self.chunk_document(text)
        
        # Generate embeddings
        chunks_with_embeddings = await self.generate_embeddings(chunks)
        
        # Add chunks to document
        for chunk in chunks_with_embeddings:
            document.add_chunk(chunk)
        
        return document
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        docx_file = io.BytesIO(content)
        doc = DocxDocument(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_markdown(self, content: bytes) -> str:
        """Extract text from Markdown"""
        # Markdown is plain text, just decode UTF-8
        return content.decode('utf-8')
    
    async def chunk_document(self, content: str, chunk_size: int = 1000) -> List[DocumentChunk]:
        """Chunk document into smaller pieces"""
        chunks = []
        words = content.split()
        current_chunk = []
        current_length = 0
        chunk_index = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            if current_length + word_length > chunk_size and current_chunk:
                # Save current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append(DocumentChunk(
                    chunk_index=chunk_index,
                    content=chunk_text
                ))
                chunk_index += 1
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-self.chunk_overlap//10:] if len(current_chunk) > self.chunk_overlap//10 else current_chunk
                current_chunk = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk)
            else:
                current_chunk.append(word)
                current_length += word_length
        
        # Add remaining chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(DocumentChunk(
                chunk_index=chunk_index,
                content=chunk_text
            ))
        
        return chunks
    
    async def generate_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Generate embeddings for document chunks"""
        texts = [chunk.content for chunk in chunks]
        embeddings = await self.embedding_service.embed_texts(texts)
        
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
        
        return chunks

